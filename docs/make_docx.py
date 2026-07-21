"""Build the Plant Brain project context document as a .docx.

Run:  py docs/make_docx.py    ->  docs/Plant_Brain_Context.docx
Diagrams come from docs/img/ (render them first with make_diagrams.py).
"""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

HERE = Path(__file__).parent
IMG = HERE / "img"
OUT = HERE / "Plant_Brain_Context.docx"

BLACK = RGBColor(0, 0, 0)
GREY = RGBColor(0x55, 0x55, 0x55)


def set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def para(doc, text="", size=10.5, bold=False, italic=False, color=BLACK,
         space_after=6, align=None, style=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if text:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.bold = bold
        r.italic = italic
        r.font.color.rgb = color
        r.font.name = "Calibri"
    return p


def rich(doc, parts, size=10.5, space_after=6):
    """parts = [(text, bold, italic), ...]"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    for text, b, i in parts:
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.bold = b
        r.italic = i
        r.font.name = "Calibri"
        r.font.color.rgb = BLACK
    return p


def h1(doc, n, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "12")
    top.set(qn("w:space"), "6")
    top.set(qn("w:color"), "000000")
    bdr.append(top)
    pPr.append(bdr)
    r = p.add_run(f"{n}   {text}")
    r.font.size = Pt(15)
    r.bold = True
    r.font.name = "Calibri"
    r.font.color.rgb = BLACK
    p.paragraph_format.keep_with_next = True
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(11.5)
    r.bold = True
    r.font.name = "Calibri"
    r.font.color.rgb = BLACK
    p.paragraph_format.keep_with_next = True
    return p


def bullets(doc, items, size=10.5):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        if isinstance(it, str):
            it = [(it, False, False)]
        for text, b, i in it:
            r = p.add_run(text)
            r.font.size = Pt(size)
            r.bold = b
            r.italic = i
            r.font.name = "Calibri"


def numbered(doc, items, size=10.5):
    for it in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        if isinstance(it, str):
            it = [(it, False, False)]
        for text, b, i in it:
            r = p.add_run(text)
            r.font.size = Pt(size)
            r.bold = b
            r.italic = i
            r.font.name = "Calibri"


def table(doc, headers, rows, widths=None, size=9.5):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        r = p.add_run(htext)
        r.bold = True
        r.font.size = Pt(size)
        r.font.name = "Calibri"
        set_cell_bg(hdr[i], "EDEDED")
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            bold = False
            text = str(val)
            if text.startswith("**") and text.endswith("**"):
                bold, text = True, text[2:-2]
            r = p.add_run(text)
            r.font.size = Pt(size)
            r.bold = bold
            r.font.name = "Calibri"
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def figure(doc, filename, caption, width=6.4):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(IMG / filename), width=Inches(width))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(12)
    r = c.add_run(caption)
    r.font.size = Pt(9)
    r.italic = True
    r.font.name = "Calibri"
    r.font.color.rgb = GREY


def feature(doc, fid, title, body, why):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"{fid}   {title}")
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.name = "Calibri"
    p.paragraph_format.keep_with_next = True
    para(doc, body, size=10.5, space_after=2)
    q = doc.add_paragraph()
    q.paragraph_format.left_indent = Inches(0.3)
    q.paragraph_format.space_after = Pt(6)
    r1 = q.add_run("Why:  ")
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.name = "Calibri"
    r2 = q.add_run(why)
    r2.font.size = Pt(10)
    r2.italic = True
    r2.font.name = "Calibri"
    r2.font.color.rgb = GREY


# ============================================================ build
doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.9)
sec.bottom_margin = Inches(0.9)
sec.left_margin = Inches(0.9)
sec.right_margin = Inches(0.9)

st = doc.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(10.5)

# ---- title page ----
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
r = p.add_run("ET AI HACKATHON 2026  ·  PROBLEM STATEMENT 8")
r.font.size = Pt(9.5)
r.bold = True
r.font.name = "Calibri"
r.font.color.rgb = GREY

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
r = p.add_run("Plant Brain")
r.font.size = Pt(30)
r.bold = True
r.font.name = "Calibri"

para(doc, "Project context document — complete build record", size=13,
     color=GREY, space_after=14)

para(doc,
     "A unified industrial knowledge platform. It ingests heterogeneous plant documents "
     "into a single knowledge graph, answers questions with citations that resolve to a "
     "page or an audio timestamp, watches equipment history for failure patterns, captures "
     "retiring operators' spoken knowledge in any language, and maps quality findings back "
     "to the standard clause that governs them. Everything runs locally — no cloud API at "
     "any point.", space_after=12)

table(doc, ["Property", "Value", "Property", "Value"], [
    ["Datastore", "Neo4j 5.26", "Reasoning model", "qwen2.5:7b-instruct"],
    ["Embeddings", "nomic-embed-text, 768 d", "Speech", "faster-whisper large-v3"],
    ["GPU footprint", "6 GB", "Codebase", "~4,200 lines, 13 commits"],
    ["Graph", "15 labels, 279 nodes", "Evaluation", "21/21, 12.0 s mean"],
], widths=[1.35, 1.95, 1.5, 1.9])

# ---- 01 ----
h1(doc, "01", "The problem")
para(doc,
     "A large industrial plant runs seven to twelve disconnected document systems. "
     "Engineering drawings in one, maintenance work orders in another, operating procedures "
     "in a third, inspection records in a fourth, regulatory paperwork buried in email. "
     "Staff spend roughly 35% of their working hours searching for information that already "
     "exists somewhere in the organisation.")
para(doc,
     "That fragmentation contributes to 18–22% of unplanned downtime, because maintenance "
     "decisions get made without complete equipment history. An engineer deciding whether a "
     "pump can wait until the next outage often cannot see that its three sister units "
     "failed the same way last year.")
para(doc,
     "Separately and more permanently: around 25% of India's experienced industrial "
     "engineers and operators retire within the next decade, taking undocumented operational "
     "knowledge with them. The knowledge that never made it into a procedure — which machine "
     "runs hot in summer, which interval the manual gets wrong — leaves with them.")

h2(doc, "What Problem Statement 8 asks for")
para(doc,
     "An AI platform that ingests heterogeneous industrial documents, extracts entities "
     "(equipment tags, process parameters, regulatory references, personnel, dates), builds "
     "a unified knowledge graph, and makes the collective intelligence queryable and "
     "actionable at the point of need.")

table(doc, ["Judging criterion", "Weight", "How this build addresses it"], [
    ["Innovation", "25%", "Multilingual voice capture of retiring operators; procedure-vs-practice divergence detection"],
    ["Business impact", "25%", "Every alert states avoidable downtime in hours, computed from that failure mode's own history"],
    ["Technical excellence", "20%", "Hybrid retrieval with graph expansion; computed confidence; measured ablation study"],
    ["Scalability", "15%", "Blackboard architecture — the eight deferred agents require zero changes to existing ones"],
    ["User experience", "15%", "Streaming answers, clickable citations to page and audio timestamp, live graph explorer"],
], widths=[1.4, 0.7, 4.6])

# ---- 02 ----
h1(doc, "02", "What we built")
para(doc,
     "A system that ingests factory documents into a Neo4j knowledge graph, answers "
     "questions with citations that resolve to a specific page or audio second, watches "
     "equipment history for failure patterns without being asked, captures retiring "
     "operators' spoken knowledge in any language, and maps quality findings back to the "
     "standard clause that governs them.")

h2(doc, "The two differentiators")
para(doc, "Most of the system is substrate. Two capabilities make it distinct, and every "
          "scope decision protected them:")
numbered(doc, [
    [("Voice capture", True, False),
     (" — multilingual audio from retiring staff, translated to English and structured into "
      "typed knowledge claims, with the original audio playable at the exact second the claim "
      "was spoken.", False, False)],
    [("Procedure-vs-practice divergence", True, False),
     (" — detecting where the written procedure and what operators actually do have drifted "
      "apart. No document system surfaces this, because it requires holding unverified spoken "
      "practice and approved policy as separate things and then deliberately comparing them.",
      False, False)],
])
rich(doc, [
    ("The insight that ties them together. ", True, False),
    ("The retiring-workforce problem and the fragmented-documents problem are usually treated "
     "separately. They are the same problem: knowledge that exists but is not reachable at the "
     "moment of decision. Divergence detection is what happens when you make both reachable in "
     "one graph — the procedure says monthly, three operators independently say fortnightly "
     "during monsoon, and the system can now notice.", False, False),
], space_after=10)

# ---- 03 ----
h1(doc, "03", "Scope decisions")
para(doc,
     "What was deliberately not built matters as much as what was. Each cut below was a full "
     "day of work carrying real risk, and none adds a capability the platform does not already "
     "demonstrate through another door.")
table(doc, ["Not built", "Reason"], [
    ["P&ID / drawing parsing (computer vision)", "A full day of work, high risk, and every competing team will demo it"],
    ["OCR / scanned documents", "Adds a failure mode; the corpus is generated, so we control formats"],
    ["Email ingestion", "Another door into the same pipeline, no new capability"],
    ["Text-to-Cypher / counting questions", "Not in the demo script"],
    ["Alert subgraph image rendering", "The evidence chain shows as a structured list instead"],
    ["Mobile app / on-device inference", "Server-hosted responsive web only"],
    ["Any additional AI model", "Headroom exists — it was spent on stability instead"],
], widths=[2.6, 4.1])
para(doc,
     "Both computer vision and OCR appear in the problem statement's suggested technologies. "
     "They were cut consciously, not overlooked — see section 18 for how they fit the "
     "production design.")

# ---- 04 ----
h1(doc, "04", "Platform and models")
para(doc,
     "Hard constraint: a single laptop with a 6 GB GPU. The original plan assumed 8 GB, so "
     "allocation was re-benchmarked rather than assumed — a decision that changed which model "
     "runs in production.")
table(doc, ["Model", "Placement", "Size", "Role"], [
    ["qwen2.5:7b-instruct", "82% GPU / 18% CPU", "4.7 GB", "Relation binding, transcript structuring, answer synthesis, alert narratives"],
    ["nomic-embed-text", "GPU", "275 MB", "All embeddings, 768 dimensions — locked on hour one"],
    ["faster-whisper large-v3", "CPU, int8", "1.5 GB", "Speech to English in one pass"],
], widths=[1.5, 1.25, 0.7, 3.25])

h2(doc, "Decisions worth defending")
rich(doc, [("The 7B stayed despite not fitting in VRAM. ", True, False),
           ("Benchmarking showed it splits 82/18 GPU/CPU and sustains a stable 22–28 tokens "
            "per second. A 4B alternative fit entirely in VRAM but showed severe variance — "
            "71 tok/s on one call, then 4.9 on the next. Stable and slower beats fast and "
            "unpredictable when there is one take on stage. The 4B remains as emergency "
            "fallback.", False, False)])
rich(doc, [("Whisper is on CPU deliberately. ", True, False),
           ("On GPU it would contend with the 7B for VRAM. At int8 on CPU it runs several "
            "times faster than realtime, and transcription happens once, offline, before the "
            "demo. Downsizing to a smaller Whisper was rejected outright — Indian-language "
            "accuracy degrades sharply.", False, False)])
rich(doc, [("Embedding dimension is irreversible. ", True, False),
           ("768 was fixed in the first thirty minutes and never revisited. Changing the "
            "embedding model later invalidates every vector index and forces a complete "
            "re-ingest.", False, False)])

h2(doc, "Stack")
bullets(doc, [
    [("Neo4j 5.26", True, False), (" — single datastore. Native vector index plus Lucene "
     "full-text means no separate vector database to keep in sync.", False, False)],
    [("Ollama", True, False), (" — all inference, fully local.", False, False)],
    [("FastAPI", True, False), (" — backend, with server-sent events for streaming.", False, False)],
    [("Plain responsive HTML/JS", True, False), (" — no framework, no build tooling.", False, False)],
    [("faster-whisper", True, False), (" — speech recognition. ", False, False),
     ("APScheduler", True, False), (" — the Watch agent's timer.", False, False)],
])

# ---- 05 ----
h1(doc, "05", "The corpus")
para(doc,
     "Every document is produced by a seeded, deterministic generator — nothing is "
     "hand-authored. When the schema changes you regenerate rather than rewrite. Critically, "
     "the generators assert their own output: every planted pattern is verified on every run, "
     "and the script exits non-zero if any is broken.")
table(doc, ["Artefact", "Count", "Detail"], [
    ["Equipment register", "12", "6 pumps sharing one class (needed for fleet analysis), 3 exchangers, 3 control valves"],
    ["Work orders", "119", "61 preventive · 40 corrective · 18 inspection, spanning three years"],
    ["Inspection records", "6", "2 deliberately expired"],
    ["Procedures (SOP PDFs)", "8", "826–910 words each"],
    ["Quality standard", "1", "QS-001, 9 clauses, 835 words, original text"],
    ["Quality KPIs", "72", "3 processes × 2 metrics × 12 months"],
    ["Non-conformance records", "6", "each linked to equipment and a clause"],
    ["Voice clips", "3", "Tamil, Hindi, one more — to be recorded"],
], widths=[1.7, 0.6, 4.4])

rich(doc, [("Why procedures are 800+ words. ", True, False),
           ("Short documents produce a single chunk each, which makes the retrieval demo look "
            "trivial — there is nothing to retrieve from. The generator enforces the word "
            "count as a verified assertion.", False, False)])

h2(doc, "The six planted patterns")
para(doc, "These are not decoration. Every demo moment depends on one of them existing, so "
          "each is asserted by the generator.")
table(doc, ["#", "Pattern", "Construction", "What it demonstrates"], [
    ["P1", "Overdue", "P-101A, four seal failures at 88/84/86-day intervals, last one 80 days before demo", "MTBF ≈ 86 d, risk ratio ≈ 0.93 — clears the 0.8 alert threshold"],
    ["P2", "Sibling exposure", "Plugged-strainer failures on P-101A/B/C only; P-102A, P-102B, P-103A share the class but have no such history", "One failure identifies five sister pumps — two with confirmed history, three needing preventive checks"],
    ["P3", "Chronic", "P-102A, four vibration failures inside twelve months", "Same fault repeatedly — nobody fixed the root cause"],
    ["P4", "**Divergence**", "SOP-114 states monthly strainer cleaning; three voice transcripts independently say fortnightly during monsoon", "The winning finding. Written policy vs. actual practice"],
    ["Q1", "Capability drift", "PRC-301 process capability falling 1.41 → 1.34 → 1.21", "Breaches the 1.33 minimum; ties to CV-301A's expired calibration record"],
    ["Q2", "Defect trend", "PRC-101 defect rate rising 165 → 240 → 310 → 420 through monsoon months", "Three consecutive rises; ties to the strainer story and clause 8.5.1"],
], widths=[0.35, 1.15, 2.7, 2.5])

h2(doc, "Injected messiness")
para(doc,
     "Deliberate, and demonstrable — this is what proves entity resolution actually works. "
     "The same pump is written four different ways across the corpus: P-101A in work orders, "
     "Pump 101A in procedures, P101-A in one inspection record, and p-101a in a voice "
     "transcript. Date formats vary between files, and one failure code carries a deliberate "
     "typo that the pipeline must repair.")

# ---- 06 ----
doc.add_page_break()
h1(doc, "06", "System architecture")
para(doc,
     "Four agents operate on a blackboard: no agent ever calls another. They read and write "
     "shared graph state, which is why the eight deferred agents can be added without "
     "modifying a single existing one.")
figure(doc, "fig1_architecture.png",
       "Figure 1 — System architecture. Blackboard pattern: agents share graph state and never call each other.")

# ---- 07 ----
h1(doc, "07", "Ingestion pipeline")
para(doc,
     "The most consequential structural decision in the build: clean data never touches the "
     "language model. Spreadsheets are already typed; running them through an extraction "
     "model only introduces hallucination into the most reliable source available.")
para(doc,
     "The payoff is large. Structured-first means the 7B processes roughly 40 narrative "
     "chunks instead of about 500 — extraction drops from about 50 minutes to under 5, and "
     "the model's weaker relation-binding stops being a risk to the whole graph.")
figure(doc, "fig2_ingestion.png",
       "Figure 2 — Ingestion, four paths. Regex runs before the LLM; entity resolution is the convergence point.")
rich(doc, [("Entity resolution is the single point of failure. ", True, False),
           ("Three nodes for one pump means every multi-hop query silently returns nothing — "
            "not an error, just emptiness. It is verified by hand before anything downstream "
            "is trusted, and the interface exposes it as a panel so the merge is visible "
            "rather than assumed.", False, False)])

# ---- 08 ----
doc.add_page_break()
h1(doc, "08", "Graph schema")
para(doc,
     "One database, two deliberately different storage modes. Document content carries 768-d "
     "embeddings for semantic retrieval. Facts and measurements are typed nodes queried by "
     "Cypher and computed over — never embedded, because a vector search cannot do arithmetic.")
figure(doc, "fig3a_schema_core.png",
       "Figure 3a — Maintenance core schema. Equipment is the hub every other label reaches.", width=6.2)
figure(doc, "fig3b_schema_qms.png",
       "Figure 3b — Quality management extension. Same graph, same hub.", width=6.2)

rich(doc, [("The one schema rule that cannot be broken. ", True, False),
           (":TacitKnowledge and :Procedure are separate labels. Spoken operator experience is "
            "one person's unverified practice; a procedure is approved policy. Merging them "
            "would be unsafe in a plant context — and keeping them apart is the only reason "
            "divergence detection is possible at all.", False, False)])

h2(doc, "Live inventory")
table(doc, ["Node label", "Count", "Node label", "Count"], [
    ["WorkOrder", "119", "Clause", "9"],
    ["QualityKPI", "72", "Document", "8"],
    ["Chunk", "24", "Procedure", "8"],
    ["Equipment", "12", "FailureMode", "7"],
    ["Alert", "7", "NCR", "6"],
    ["InspectionRecord", "6", "Process", "3"],
    ["Area", "2", "Site / Standard", "1 / 1"],
    ["TacitKnowledge", "0", "awaiting voice recordings", ""],
], widths=[1.8, 0.8, 1.8, 0.8])

# ---- 09 ----
h1(doc, "09", "Retrieval design")
para(doc,
     "Both retrieval types are required because they fail on opposite question shapes. "
     "Keyword search catches exact tags and standard codes that embeddings miss. Dense "
     "retrieval catches paraphrase — “won't build pressure” matching “insufficient "
     "discharge head”. Neither reaches a question whose answer shares no vocabulary with it: "
     "a report on P-101C's failure never contains the string “P-101A”. Only walking "
     "equipment → same class → other equipment does.")
figure(doc, "fig4_retrieval.png",
       "Figure 4 — Ask agent. Both legs always run; graph expansion executes in the same Cypher round trip.",
       width=6.0)

h2(doc, "Blend, never branch")
para(doc,
     "The intent classifier adjusts how much text versus graph evidence is assembled; it "
     "never switches one off. Misclassifying an exclusive branch returns nothing. A wrong "
     "blend still returns something useful.")

h2(doc, "Confidence is computed, never self-reported")
para(doc,
     "A model asked to rate its own certainty produces a number with no grounding. Here the "
     "score is arithmetic over observable signals — normalised top retrieval score, the "
     "fraction of answer sentences that carry a citation, and agreement between the keyword "
     "and dense result sets. Below threshold the system abstains. In a safety context, "
     "“not present in the corpus” is the correct answer, and a confident wrong answer "
     "about plant equipment is worse than no answer at all.")

# ---- 10 ----
doc.add_page_break()
h1(doc, "10", "Divergence detection")
para(doc,
     "The differentiating finding, and the hardest logic in the build. Naive approaches fail "
     "in two specific ways, both of which had to be designed around.")
figure(doc, "fig5_divergence.png",
       "Figure 5 — Divergence detection. Two gates before comparison; frequency set-difference, not intersection.",
       width=6.0)

h2(doc, "Failure mode 1 — embedding similarity alone")
para(doc,
     "On a small corpus, semantic similarity fires constantly on unrelated documents. Two "
     "texts about plant maintenance look similar to a vector regardless of whether they "
     "concern the same task. The fix is a same-topic gate: the procedure and the claim must "
     "apply to the same equipment and share an explicit maintenance-activity keyword before "
     "their frequencies are compared at all.")

h2(doc, "Failure mode 2 — the intersection trap")
para(doc,
     "Operators almost always cite the official interval while rejecting it: “the book says "
     "monthly, but we do it fortnightly.” A naive set intersection sees “monthly” in both "
     "texts and concludes the two agree — silently missing the exact finding it was built to "
     "catch. The detector instead isolates the frequency the operator states that the "
     "procedure does not, which is the actual deviating practice.")
para(doc,
     "A third refinement: findings collapse to one alert per procedure rather than one per "
     "matched keyword, which otherwise produced three near-duplicate alerts for the same "
     "underlying divergence.")

# ---- 11 ----
h1(doc, "11", "QMS integration")
para(doc,
     "Quality management is a specialization of the same architecture, not a second system — "
     "same graph, same agents, same retrieval logic. Only the ingested standards pack and the "
     "KPI definitions differ.")
h2(doc, "Clause-aware chunking")
para(doc,
     "The standard splits on clause numbers, not fixed token windows. A clause is the natural "
     "retrieval unit in compliance; splitting one mid-sentence produces answers that cite half "
     "a requirement, which is worse than useless in an audit.")
h2(doc, "Numbers are never embedded")
para(doc,
     "This is where pure-RAG architectures typically fail. Capability and defect rates load "
     "into typed :QualityKPI rows an agent queries and computes over. Embedding “Cpk = 1.21” "
     "as a sentence and hoping a vector search retrieves it correctly cannot support "
     "arithmetic, threshold tests, or trend detection. Thresholds come from configuration, so "
     "every computed figure is reproducible and carries the exact query that produced it.")
h2(doc, "Compliance mapping")
para(doc,
     "Non-conformance records link equipment to the clause they breach. That is what lets "
     "“why is PPM rising on the feed transfer process” traverse from a number, to the "
     "process, to the equipment, to the open NCR, to governing clause 8.5.1 — and then to the "
     "monsoon strainer story that explains it. No vector search reconstructs that chain.")
h2(doc, "Industry-agnostic by configuration")
para(doc,
     "The pipeline is fixed; the standards pack swaps. Configuration defines which base "
     "standard applies, its clause taxonomy, and the industry's KPI definitions — defect rate "
     "for automotive, batch yield for pharma, lot traceability for food. Graph schema and "
     "agent logic stay identical. That is what makes this one architecture rather than one "
     "build per industry.")

# ---- 12 ----
doc.add_page_break()
h1(doc, "12", "Feature inventory")
para(doc,
     "Twenty-six implemented capabilities. The note under each states why it exists — in most "
     "cases the design was forced by a specific failure it prevents.")

h2(doc, "Knowledge construction")
feature(doc, "A1", "Deterministic, self-verifying corpus generator",
        "Every document is produced by a seeded script that asserts its own output; planted patterns are verified on every run.",
        "Regenerate rather than rewrite when the schema changes, and a corpus that verifies itself cannot silently lose the patterns the demo depends on.")
feature(doc, "A2", "Entity resolution across four spellings",
        "The same pump appears four ways across the corpus; all canonicalise to one node with an alias list.",
        "The single point of failure. Three nodes for one pump means every multi-hop query silently returns nothing.")
feature(doc, "A3", "Structured-first ingestion",
        "Spreadsheet data becomes typed nodes with no LLM in the path.",
        "Running typed data through an extraction model introduces hallucination into the cleanest source available — and it cuts model workload by over 90%.")
feature(doc, "A4", "Regex-before-LLM relation binding",
        "Regex finds entities first; the LLM only binds relations between entities already found, against a closed vocabulary.",
        "A 7B model is weak at open-ended extraction but adequate at choosing among known candidates. This plays to that boundary rather than across it.")
feature(doc, "A5", "Failure-code repair",
        "A mistyped ISO 14224 code is mapped back to the closed vocabulary on load, raw value preserved.",
        "Real maintenance records contain typos, and a closed vocabulary is what makes the analytics work at all.")
feature(doc, "A6", "Clause-aware chunking",
        "The standard splits on clause boundaries rather than token counts.",
        "A compliance answer citing half a requirement is worse than no answer.")
feature(doc, "A7", "Structured quality store — numbers never embedded",
        "Capability and defect values are typed rows computed over, not text in a vector index.",
        "Embeddings cannot do arithmetic or trend detection. This is the specific thing pure-RAG quality systems get wrong.")

h2(doc, "Answering")
feature(doc, "B1", "Hybrid retrieval with RRF fusion",
        "Keyword and dense retrieval run in parallel on every question and fuse by reciprocal rank.",
        "They fail on opposite question shapes. Running only one guarantees a class of question returns nothing.")
feature(doc, "B2", "Graph expansion in a single round trip",
        "Two hops — equipment to work orders to failure modes to siblings to operator knowledge — inside the same Cypher statement.",
        "It answers questions whose answer shares no vocabulary with the question, and one statement keeps latency flat.")
feature(doc, "B3", "Multi-form anchoring",
        "Anchors on an explicit tag, a class or area phrase, a process name, or a clause number.",
        "People rarely quote asset tags. Without phrase anchoring, questions naming a class wrongly abstained — found in testing, fixed.")
feature(doc, "B4", "Blend ratio, never an exclusive branch",
        "Intent adjusts the evidence mix; it never switches a retrieval leg off.",
        "Misclassifying an exclusive branch returns nothing; a wrong blend still returns something useful.")
feature(doc, "B5", "Symptom-aware evidence ordering",
        "Question wording maps to ISO failure modes; matching history rows sort to the top.",
        "A real failure — the relevant vibration record sat seventh of seven, the model overlooked it, and the system abstained on a question it could answer.")
feature(doc, "B6", "Analytics leg for numeric questions",
        "Capability and trend questions attach computed KPI evidence — series, threshold test, explicit verdict — ahead of document text.",
        "Two further real failures. Buried mid-prompt the model hedged around its own computed answer; and when the block said only PRC-301 it could not bind that to a question naming the feed flow control process.")
feature(doc, "B7", "Computed confidence with abstention",
        "Arithmetic over retrieval strength, citation density and result-set agreement; below threshold, declines to answer.",
        "Self-reported confidence is ungrounded, and in a plant context a confident wrong answer is more dangerous than none.")
feature(doc, "B8", "Resolvable citations",
        "Markers resolve to a document page, an audio second, or the exact query run against the structured store — and open on click.",
        "An uncited reference is worse than no answer in compliance, and a citation you cannot open is just a claim.")
feature(doc, "B9", "Streaming responses with model pre-warm",
        "Answers stream token by token; models load at startup.",
        "First tokens appear in about two seconds instead of a blank 10–20 second wait, and pre-warming removes the cold-start stall entirely.")
feature(doc, "B10", "Show-the-query toggle",
        "The Cypher behind any answer is inspectable from the answer itself.",
        "Makes the system auditable rather than magic — the difference between a demo and a tool an engineer would trust.")

h2(doc, "Proactive intelligence")
feature(doc, "C1", "Reliability detection — overdue, chronic, sibling exposure",
        "MTBF and risk ratio from failure intervals; repeat-mode counting; same-class propagation on failure.",
        "All three are arithmetic, so a finding can be defended line by line. Sibling exposure turns one pump's failure into fleet-wide preventive action.")
feature(doc, "C2", "Procedure-vs-practice divergence (needs recordings)",
        "Detects drift between written SOP and actual practice, gated on shared equipment and shared activity before any comparison.",
        "The finding no document system surfaces. Proven against synthetic claims; awaiting real recordings.")
feature(doc, "C3", "Quality detection — capability and defect trend",
        "Capability below the configured minimum; defect rate rising three consecutive months.",
        "Computed from the structured store rather than retrieved, so the verdict is a calculation, not a paraphrase.")
feature(doc, "C4", "Business impact in hours",
        "Each alert states avoidable downtime derived from that failure mode's history on that equipment class.",
        "Turns a technical finding into a number a plant manager acts on. Narratives are constrained to facts in the data after the model invented a cost figure during testing.")
feature(doc, "C5", "Evidence chains on every alert",
        "Alerts link back to the specific work orders, chunks and claims that produced them.",
        "A finding you cannot trace is a finding nobody will act on.")
feature(doc, "C6", "Compliance mapping",
        "Findings map to the governing clause through the non-conformance register.",
        "An auditor's first question is which requirement was breached — and that traversal is graph-shaped.")

h2(doc, "Interface and evidence of correctness")
feature(doc, "D1", "Knowledge graph explorer",
        "Interactive force-directed view, architecture overview plus per-equipment focus. No external library.",
        "A CDN-hosted graph library fails on an air-gapped demo laptop. The focus view doubles as visual proof that entity resolution worked.")
feature(doc, "D2", "Entity resolution panel",
        "Shows every merged spelling and what it resolved to.",
        "Makes the least visible and most important part of the pipeline demonstrable at a glance.")
feature(doc, "D3", "Stopwatch and detected-language badge",
        "Live timing on every question; source language shown on spoken claims.",
        "Both are explicit judging criteria, and both expose something already computed rather than adding a failure mode.")
feature(doc, "D4", "Evaluation harness with retrieval ablation",
        "Twenty-one questions scored across four retrieval configurations, producing a chart.",
        "Converts “our retrieval design is good” into a measurement, and isolates exactly what the graph contributes.")
feature(doc, "D5", "Configurable standards pack",
        "Standard identity, clause taxonomy and KPI thresholds live in configuration.",
        "The difference between one architecture and one build per industry.")

# ---- 13 ----
doc.add_page_break()
h1(doc, "13", "Build record")
para(doc,
     "Thirteen commits across two working days. The rollback point v0.1-day2-stable is "
     "tagged, and a database snapshot exists at backups/neo4j.dump.")
table(doc, ["Commit", "Date", "What landed"], [
    ["e3dcd12", "20 Jul", "**Tagged stable.** Corpus generator, graph pipeline, Ask + Watch agents, web UI"],
    ["78b0df8", "20 Jul", "Streaming answers, evidence-quality fixes, eval harness"],
    ["babbec3", "20 Jul", "Symptom-mode row ordering, graph-first evidence"],
    ["df4e56b", "20 Jul", "Class/area anchors, inspection evidence, confidence-gate fix"],
    ["3fbe705", "20 Jul", "Baseline eval and ablation results"],
    ["a2f3522", "20 Jul", "Post-fix eval, regression-guard questions"],
    ["af54ced", "20 Jul", "**QMS integration** — clause-aware standard, structured KPIs, retrieval router"],
    ["5da0721", "20 Jul", "Analytics binding — process names in KPI evidence, computed-first order"],
    ["94540fb", "20 Jul", "QMS alert cards, README architecture note, QMS eval questions"],
    ["c1e3cf4", "21 Jul", "**Knowledge graph visualization**"],
    ["5b7dcfd", "21 Jul", "VSCode launch config"],
    ["6e692e4", "21 Jul", "Abstain-detection fix — eval reaches 21/21"],
    ["77e190c", "21 Jul", "Architecture document"],
], widths=[0.9, 0.7, 5.1])

h2(doc, "Codebase")
table(doc, ["Area", "Files", "Purpose"], [
    ["Generation", "datagen/ — 3 files", "Maintenance corpus, SOP text, QMS corpus"],
    ["Pipeline", "pipeline/ — 9 files", "Config, schema, entity resolution, four loaders, Ask, Watch"],
    ["Application", "app.py + static/ — 4 files", "FastAPI backend, chat UI, graph explorer, eval chart"],
    ["Verification", "scripts/eval.py", "21-question harness with four-way ablation"],
    ["Documentation", "docs/ — 4 files", "Architecture, this document, diagrams, voice recording guide"],
], widths=[1.2, 1.9, 3.6])

# ---- 14 ----
h1(doc, "14", "Engineering log")
para(doc,
     "Every defect below was found by running the system rather than reading it, and each "
     "drove a design change that survives in the code. They are recorded because how a system "
     "was debugged is stronger evidence of its soundness than a claim that it works.")

log = [
    ("Truncated evidence cut the answer out of the evidence",
     "Chunk evidence was capped at 1,200 characters — which sliced off SOP-114's monthly-interval sentence, the exact text the flagship strainer question needed. The system abstained on a question it had the answer to. Fixed by passing full chunks and raising the model's context window."),
    ("Citation markers collided with the documents' own numbering",
     "Answers cited [1], but the SOPs contain numbered sections, so the model began citing section numbers as if they were evidence. Switched to unambiguous [E1] markers."),
    ("A relevant record at the bottom of a table was invisible",
     "The flagship vibration question failed because the one matching record sat seventh of seven rows. The model summarised the visible rows and concluded no such history existed. Question wording now maps to failure modes and matching rows sort to the top."),
    ("Questions naming a class instead of a tag anchored nothing",
     "“The control valves in the process area” retrieved procedure text only and wrongly abstained. Class and area phrases now resolve to equipment sets — and correctly exclude valves in other areas."),
    ("The confidence gate was overwriting correct answers",
     "A substantive, correctly cited answer could be replaced wholesale with “not present in the corpus” because a structurally noisy overlap term dragged the score down. Weights recalibrated toward retrieval strength and citation density."),
    ("The model invented a cost figure",
     "An alert narrative asserted “$10,000 per hour of lost production” — a number that appears nowhere in the data. Narrative generation is now explicitly constrained to facts present in the alert."),
    ("Divergence detection concluded agreement on the divergence",
     "Operators say “the book says monthly, we do fortnightly” — so “monthly” appears in both texts and set intersection declared them in agreement, silently missing the finding. Now isolates the frequency the operator states that the procedure does not."),
    ("Computed answers hedged when buried mid-prompt",
     "The capability question kept refusing despite a clear NOT CAPABLE flag in evidence. Two causes: computed blocks sat below document text, and the block named only PRC-301 while the question said “the feed flow control process”. Computed evidence now leads, and carries the process name."),
    ("The system did not recognise its own refusal",
     "The model writes “Not present in the corpus [E1][E2].” — citation markers land before the period, so a period-inclusive string match failed. The refusal was correct; the detection of it was not. Matching now ignores trailing punctuation."),
    ("A planted pattern was contaminated by random data",
     "Randomly generated corrective work orders happened to give P-102A a plugged-strainer failure, which broke the sibling-exposure pattern's clean split. Protected combinations are now excluded from random generation."),
]
for title, body in log:
    rich(doc, [("• " + title + ".  ", True, False), (body, False, False)], space_after=7)

h2(doc, "Environment findings")
bullets(doc, [
    "The GPU is 6 GB, not the 8 GB assumed — model allocation was re-benchmarked rather than inherited.",
    "Neo4j runs as a foreground console process, so it stops when its terminal closes. Installing it as a service is recommended before demo day.",
    "Machine sleep evicts Ollama models mid-inference, surfacing as a server error; the next call reloads them.",
    "Python must be invoked as py — plain python hits the Windows Store stub.",
])

# ---- 15 ----
doc.add_page_break()
h1(doc, "15", "Measured results")
para(doc,
     "A question passes only if the abstain decision is correct and, for answerable questions, "
     "an expected source appears in the citations the answer actually used. Citing the right "
     "document by luck while answering from elsewhere does not pass.")
table(doc, ["Suite", "Passed", "Mean latency", "Composition"], [
    ["Full pipeline", "21 / 21", "12.0 s", "12 equipment/procedure · 3 QMS router paths · 3 abstain traps · 3 regression guards"],
], widths=[1.2, 0.8, 1.0, 3.7])

h2(doc, "Retrieval ablation")
para(doc, "Measured on the 15-question maintenance suite, disabling one retrieval leg at a time:")
table(doc, ["Configuration", "Passed", "Failure character"], [
    ["Full — hybrid + graph", "15 / 15", "—"],
    ["Dense only (no keyword)", "14 / 15", "an exact-threshold lookup"],
    ["Keyword only (no dense)", "14 / 15", "same"],
    ["**No graph expansion**", "**11 / 15**", "**every failure is an equipment-history question**"],
], widths=[2.0, 0.9, 3.8])

para(doc,
     "The ablation makes the architectural argument precisely: removing graph traversal fails "
     "exactly the questions whose answers share no vocabulary with the question. That is the "
     "class of question a vector database alone cannot serve, and it is the reason this is a "
     "knowledge graph rather than a document index.")
rich(doc, [("Honest reading of the single-leg rows. ", True, False),
           ("Dense-only and keyword-only score higher than they would in a pure-RAG system, "
            "because graph anchoring still carries them — the text legs are crippled but "
            "equipment traversal is not. The clean, defensible finding is the no-graph row. "
            "Lead with that one.", False, False)])

# ---- 16 ----
h1(doc, "16", "Running the system")
para(doc,
     "Two background services must be up before the application starts; the graph itself "
     "persists on disk, so the ingestion pipeline is a one-time operation, not part of "
     "routine startup.")
h2(doc, "Routine startup")
numbered(doc, [
    [("Neo4j", True, False), (" — neo4j.bat console with JAVA_HOME set. Leave the window "
     "open. Check first with a port test on 7687: if it already answers, skip this step, "
     "because a second instance will fail on the database lock.", False, False)],
    [("Ollama", True, False), (" — usually already running; otherwise ollama serve.", False, False)],
    [("Application", True, False), (" — py app.py, or F5 in VSCode. Serves on port 8000.", False, False)],
])
h2(doc, "Full rebuild from scratch")
para(doc, "Only needed after regenerating the corpus or wiping the database:")
numbered(doc, [
    "py datagen/generate_corpus.py  and  py datagen/generate_qms.py",
    "py pipeline/schema.py — constraints and indexes, before any ingestion",
    "py pipeline/load_structured.py, then ingest_narrative.py, then load_qms.py",
    "py pipeline/voice_capture.py once audio clips are in place",
])
para(doc, "Verification: py scripts/eval.py for the full ablation, or --quick for the "
          "single-pass suite.")

# ---- 17 ----
h1(doc, "17", "Demo script")
para(doc, "Three minutes. Steps 4 and 6 are what win — they should be rehearsed hardest.")
numbered(doc, [
    [("Run the ingest", True, False), (" — the graph counter ticks up live.", False, False)],
    [("Entity resolution panel", True, False), (" — four spellings visibly merged to one asset.", False, False)],
    [("Ask the vibration question", True, False), (" — “P-101A keeps tripping on high "
     "vibration, has this happened before?” Cited answer, stopwatch showing seconds, click a "
     "citation to land on the exact page.", False, False)],
    [("Click a different citation", True, False), (" — Tamil audio plays, with the English "
     "claim beside it.", False, False)],
    [("Alert panel", True, False), (" — five sister pumps identified, evidence chain shown, "
     "downtime-avoided figure stated.", False, False)],
    [("Divergence panel", True, False), (" — “SOP-114 says monthly. Three operators say "
     "fortnightly during monsoon.”", False, False)],
    [("Ask something outside the corpus", True, False), (" — the system abstains.", False, False)],
])
rich(doc, [("Record the video before polishing. ", True, False),
           ("Single laptop, live system, 6 GB GPU — thirty minutes of recording is complete "
            "insurance against anything failing on stage.", False, False)])

# ---- 18 ----
h1(doc, "18", "Status and scale")
table(doc, ["Component", "Status"], [
    ["Corpus generation, graph build, entity resolution", "Operational"],
    ["Ask agent — retrieval, synthesis, citations, abstention", "Operational"],
    ["Watch agent — 5 of 6 detectors firing", "Operational"],
    ["QMS — clauses, KPIs, compliance mapping", "Operational"],
    ["Web interface, graph explorer, eval harness", "Operational"],
    ["Voice capture — code complete, pipeline proven", "**Awaiting 3 recordings**"],
    ["Divergence detection — logic proven on synthetic claims", "**Unblocks with voice**"],
    ["Telegram delivery — stretch item", "Deferred"],
], widths=[4.4, 2.3])

h2(doc, "The one open dependency")
para(doc,
     "Three voice recordings of roughly ninety seconds each, in three languages, spoken "
     "naturally rather than read. Each must independently state that P-101 strainers need "
     "fortnightly cleaning during monsoon while the manual says monthly. Everything "
     "downstream — transcription, translation, claim extraction, divergence detection — is "
     "built and proven; it is waiting only on the audio.")

h2(doc, "Why four agents, and what the other eight are")
para(doc,
     "The production design is twelve agents across four tiers — ingestion, query, "
     "intelligence, delivery. The prototype implements four. The remaining eight are email "
     "triage, Cypher generation, root cause analysis, compliance auditing, lessons-learned "
     "clustering, multi-channel notification, drawing parsing, and OCR.")
para(doc,
     "Because the architecture is a blackboard — agents read and write shared graph state and "
     "never call each other — adding the missing eight requires changing zero existing agents. "
     "A new agent subscribes to the graph, writes its findings back as nodes, and every "
     "existing agent can consume them without knowing it exists.")
para(doc,
     "Knowing precisely what was deferred, and why, is engineering judgment rather than "
     "omission.")

doc.save(OUT)
print("wrote", OUT)
print("paragraphs:", len(doc.paragraphs), "| tables:", len(doc.tables))
